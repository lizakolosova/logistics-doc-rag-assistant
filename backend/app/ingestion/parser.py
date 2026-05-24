import logging
import time
from pathlib import Path
from uuid import UUID

import fitz  # PyMuPDF
from docx import Document

from app.config import settings
from app.exceptions import (
    DocumentParseError,
    DocumentTooLargeError,
    UnsupportedFormatError,
)
from app.models.schemas import ParsedSection

logger = logging.getLogger(__name__)

_HEADING_STYLES: frozenset[str] = frozenset(f"Heading {i}" for i in range(1, 10))


def parse_pdf(file_path: Path, document_id: UUID, source_filename: str | None = None) -> list[ParsedSection]:
    """Extract text from a text-based PDF, one ParsedSection per page.

    Args:
        file_path: Path to the PDF file.
        document_id: UUID to attach to every returned section.

    Returns:
        List of ParsedSection objects, one per non-empty page.

    Raises:
        DocumentParseError: If a page is image-only or the document has no text.
    """
    display_name = source_filename or file_path.name
    start = time.monotonic()
    doc = fitz.open(str(file_path))
    sections: list[ParsedSection] = []

    try:
        for page_index in range(len(doc)):
            page = doc[page_index]
            text = page.get_text().strip()

            if not text:
                if page.get_images():
                    raise DocumentParseError(filename=display_name,
                        reason=(f"Page {page_index + 1} appears to be a scanned image. ""This system requires text-based PDFs. "
                            "Please upload an OCR'd version."))
                continue

            sections.append(ParsedSection(document_id=document_id, source_file=display_name,
                    page_number=page_index + 1, text=text, section_index=len(sections)))
    finally:
        doc.close()

    if not sections:
        raise DocumentParseError(filename=display_name, reason="Document yields no extractable text.")

    elapsed_ms = int((time.monotonic() - start) * 1000)
    total_chars = sum(len(s.text) for s in sections)
    logger.info("Parsed PDF '%s': %d pages, %d chars extracted in %dms", display_name,
        len(sections), total_chars, elapsed_ms)
    return sections


def parse_docx(file_path: Path, document_id: UUID, source_filename: str | None = None) -> list[ParsedSection]:
    """Extract text from a DOCX file, grouping paragraphs under their nearest heading.

    Args:
        file_path: Path to the DOCX file.
        document_id: UUID to attach to every returned section.

    Returns:
        List of ParsedSection objects, one per heading group (or one global section
        when no headings are present).

    Raises:
        DocumentParseError: If the document yields no extractable text.
    """
    display_name = source_filename or file_path.name
    start = time.monotonic()
    doc = Document(str(file_path))
    sections: list[ParsedSection] = []
    current_parts: list[str] = []
    current_heading: str | None = None

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if para.style and para.style.name in _HEADING_STYLES and stripped:
            if current_parts:
                sections.append(ParsedSection(document_id=document_id, source_file=display_name,
                        page_number=None, section_header=current_heading, text="\n".join(current_parts), section_index=len(sections)))
            current_parts = [stripped]
            current_heading = stripped
        elif stripped:
            current_parts.append(stripped)

    if current_parts:
        sections.append(ParsedSection(document_id=document_id, source_file=display_name,
                                      page_number=None, section_header=current_heading, text="\n".join(current_parts),
                                      section_index=len(sections)))

    if not sections:
        raise DocumentParseError(filename=display_name, reason="Document yields no extractable text.")

    elapsed_ms = int((time.monotonic() - start) * 1000)
    total_chars = sum(len(s.text) for s in sections)
    logger.info("Parsed DOCX '%s': %d sections, %d chars extracted in %dms", display_name, len(sections),
        total_chars, elapsed_ms)
    return sections


def parse_document(file_path: Path, document_id: UUID, source_filename: str | None = None) -> list[ParsedSection]:
    """Validate size/format and dispatch to the appropriate parser.

    Args:
        file_path: Path to the uploaded file.
        document_id: UUID to propagate through all returned sections.

    Returns:
        Parsed sections from the document.

    Raises:
        DocumentTooLargeError: If the file exceeds the configured size or page limit.
        UnsupportedFormatError: If the file extension is not .pdf or .docx.
        DocumentParseError: If the parser cannot extract any text.
    """
    display_name = source_filename or file_path.name
    size_mb = file_path.stat().st_size / (1024 * 1024)
    if size_mb > settings.max_file_size_mb:
        raise DocumentTooLargeError(filename=display_name, limit_type="size", actual=round(size_mb, 2),
            limit=settings.max_file_size_mb)

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        doc = fitz.open(str(file_path))
        page_count = len(doc)
        doc.close()
        if page_count > settings.max_pages:
            raise DocumentTooLargeError(filename=display_name, limit_type="pages", actual=page_count, limit=settings.max_pages)
        return parse_pdf(file_path, document_id, source_filename=display_name)

    if suffix == ".docx":
        docx_tmp = Document(str(file_path))
        section_count = len(docx_tmp.paragraphs)
        if section_count > settings.max_pages:
            raise DocumentTooLargeError(
                filename=display_name,
                limit_type="sections",
                actual=section_count,
                limit=settings.max_pages,
            )
        return parse_docx(file_path, document_id, source_filename=display_name)

    raise UnsupportedFormatError(filename=display_name, detected_type=suffix)